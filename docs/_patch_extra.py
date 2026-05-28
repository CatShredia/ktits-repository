# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

path = Path(r"C:\directory-git\ktits-repository-1\docs\Uchebnaya_Praktika_Otchet_updated.docx")
doc = Document(str(path))


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def insert_before(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)


extra = {
    "Экран «Список заказов» разделён на две части. Слева — перечень заявок текущего пользователя (номер, дата, статус, наименование); под списком выводится их количество. Справа — блок «Детали заказа» с полной информацией по выбранной строке. Вверху расположены «фильтр» (все заказы, новые, текущие, выполненные, отклонённые), кнопки «Добавить заказ» и «Обновить»; внизу — действия над выбранным заказом и кнопка «Назад» для возврата в главное меню.":
    "Экран «Список заказов» разделён на две части. Слева — перечень заявок текущего пользователя (номер, дата, статус, наименование); под списком выводится их количество. Справа — блок «Детали заказа» с полной информацией по выбранной строке. Вверху — фильтр (все, новые, текущие, выполненные, отклонённые), кнопки «Добавить заказ» и «Обновить»; внизу — действия над выбранным заказом (редактирование, отмена — по статусу).",
}

for p in doc.paragraphs:
    if p.text in extra:
        set_text(p, extra[p.text])
    if p.text.strip() == "ПРИЛОЖЕНИЕ А23":
        insert_before(p, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ22")

doc.save(str(path))
print("patched")
